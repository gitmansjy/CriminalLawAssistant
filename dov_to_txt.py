# convert_docs.py
import os
from pathlib import Path
import re

# ==================== DOCX转TXT（新格式） ====================
def convert_docx_to_txt(docx_path, txt_path):
    """
    将.docx文件转换为.txt
    使用python-docx库
    """
    try:
        from docx import Document
        
        print(f"正在转换: {docx_path}")
        
        # 打开文档
        doc = Document(docx_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # 只保留非空段落
                paragraphs.append(para.text.strip())
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        # 合并所有文本
        full_text = '\n\n'.join(paragraphs)
        
        # 写入txt文件
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        
        print(f"✅ 转换成功: {txt_path}")
        return True
        
    except ImportError:
        print("❌ 请先安装 python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False


# ==================== DOC转TXT（旧格式，需要另外处理） ====================
def convert_old_doc_to_txt(doc_path, txt_path):
    """
    将旧的.doc文件转换为.txt
    方法1：使用antiword（需要安装）
    方法2：使用catdoc（需要安装）
    方法3：使用textract（综合工具）
    """
    try:
        # 方法3：使用textract（推荐，支持多种格式）
        import textract
        
        print(f"正在转换: {doc_path}")
        
        # 提取文本
        text = textract.process(doc_path).decode('utf-8')
        
        # 清理文本
        text = re.sub(r'\n{3,}', '\n\n', text)  # 合并多余空行
        
        # 写入txt文件
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        print(f"✅ 转换成功: {txt_path}")
        return True
        
    except ImportError:
        print("❌ 请先安装 textract: pip install textract")
        print("   注意：textract需要系统依赖，Windows上可能较复杂")
        return False
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False


# ==================== 批量转换文件夹中所有Word文档 ====================
def batch_convert_word_to_txt(input_dir, output_dir=None):
    """
    批量转换文件夹中的所有Word文档
    
    Args:
        input_dir: 输入文件夹路径（存放doc/docx文件）
        output_dir: 输出文件夹路径（默认同目录）
    """
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"❌ 输入文件夹不存在: {input_dir}")
        return
    
    # 如果未指定输出目录，默认为输入目录下的txt_files文件夹
    if output_dir is None:
        output_dir = input_path / "txt_files"
    else:
        output_dir = Path(output_dir)
    
    # 创建输出目录
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("=" * 60)
    print("批量Word转TXT工具")
    print("=" * 60)
    print(f"输入目录: {input_path}")
    print(f"输出目录: {output_dir}")
    print("=" * 60)
    
    # 统计
    success_count = 0
    fail_count = 0
    
    # 遍历所有文件
    for file_path in input_path.glob("*"):
        if not file_path.is_file():
            continue
        
        # 检查文件扩展名
        suffix = file_path.suffix.lower()
        if suffix not in ['.doc', '.docx']:
            continue
        
        print(f"\n处理文件: {file_path.name}")
        
        # 生成输出文件名
        output_filename = file_path.stem + ".txt"
        output_path = output_dir / output_filename
        
        # 根据扩展名选择转换方法
        if suffix == '.docx':
            success = convert_docx_to_txt(str(file_path), str(output_path))
        else:  # .doc
            success = convert_old_doc_to_txt(str(file_path), str(output_path))
        
        if success:
            success_count += 1
        else:
            fail_count += 1
    
    print("\n" + "=" * 60)
    print(f"转换完成！成功: {success_count}, 失败: {fail_count}")
    print(f"输出文件保存在: {output_dir}")


# ==================== 如果你的文件是docx格式（最简单的情况） ====================
def quick_convert_your_files():
    """
    快速转换你的两个文件（假设是docx格式）
    """
    books_dir = Path("D:/py/CriminalLawAssistant/books")
    
    # 要转换的文件列表
    files_to_convert = [
        ("刑法.doc", "刑法.txt"),
        ("刑事诉讼法.doc", "刑事诉讼法.txt")
    ]
    
    print("=" * 60)
    print("快速转换你的法律文件")
    print("=" * 60)
    
    for doc_file, txt_file in files_to_convert:
        doc_path = books_dir / doc_file
        txt_path = books_dir / txt_file
        
        if not doc_path.exists():
            print(f"⚠️ 文件不存在: {doc_path}")
            continue
        
        print(f"\n转换: {doc_file} -> {txt_file}")
        
        # 判断实际格式（可能文件名是.doc但实际是docx）
        # 尝试读取文件头判断
        with open(doc_path, 'rb') as f:
            header = f.read(4)
        
        if header.startswith(b'PK\x03\x04'):
            # 这是docx格式（ZIP文件）
            print("  检测到: docx格式")
            convert_docx_to_txt(doc_path, txt_path)
        else:
            # 这是旧版doc格式
            print("  检测到: 旧版doc格式")
            convert_old_doc_to_txt(doc_path, txt_path)


# ==================== 如果上述方法都失败，使用文本提取器 ====================
def extract_text_with_com_bridge(doc_path, txt_path):
    """
    终极方案：使用Windows COM组件（仅限Windows）
    需要安装pywin32
    """
    try:
        import win32com.client
        
        print(f"使用COM组件转换: {doc_path}")
        
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(str(doc_path.absolute()))
        doc.SaveAs(str(txt_path.absolute()), FileFormat=2)  # 2 = wdFormatText
        doc.Close()
        word.Quit()
        
        print(f"✅ 转换成功: {txt_path}")
        return True
        
    except ImportError:
        print("❌ 请安装 pywin32: pip install pywin32")
        return False
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        return False


# ==================== 主程序 ====================
if __name__ == "__main__":
    print("=" * 60)
    print("Word文档转TXT工具")
    print("=" * 60)
    print("请选择转换方式:")
    print("1. 快速转换刑法和刑事诉讼法文件")
    print("2. 批量转换books目录下所有Word文件")
    print("3. 手动指定文件转换")
    print("=" * 60)
    
    choice = input("请输入选项 (1-3): ").strip()
    
    if choice == '1':
        quick_convert_your_files()
    
    elif choice == '2':
        books_dir = "D:/py/CriminalLawAssistant/books"
        batch_convert_word_to_txt(books_dir)
    
    elif choice == '3':
        doc_file = input("请输入Word文件路径: ").strip()
        txt_file = input("请输入输出TXT路径 (直接回车使用同名文件): ").strip()
        
        doc_path = Path(doc_file)
        if not doc_path.exists():
            print("❌ 文件不存在")
        else:
            if not txt_file:
                txt_file = str(doc_path.with_suffix('.txt'))
            
            if doc_path.suffix.lower() == '.docx':
                convert_docx_to_txt(doc_file, txt_file)
            else:
                convert_old_doc_to_txt(doc_file, txt_file)
    
    else:
        print("❌ 无效选项")
    
    input("\n按回车键退出...")